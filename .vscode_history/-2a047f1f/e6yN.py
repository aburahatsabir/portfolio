import os
import sys
from pathlib import Path

# Try to import required libraries
try:
    import PyPDF2
except ImportError:
    print("Installing PyPDF2...")
    os.system(f"{sys.executable} -m pip install PyPDF2 -q")
    import PyPDF2

try:
    from docx import Document
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document

def extract_pdf_text(pdf_path):
    """Extract text from PDF file"""
    try:
        text = ""
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += f"\n--- PAGE {page_num + 1} ---\n"
                text += page.extract_text()
        return text
    except Exception as e:
        return f"ERROR reading {pdf_path}: {str(e)}"

def extract_docx_text(docx_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " | "
                text += "\n"
        return text
    except Exception as e:
        return f"ERROR reading {docx_path}: {str(e)}"

def main():
    base_path = Path(".")
    
    print("=" * 80)
    print("DOCUMENT EXTRACTION FOR AUSTRALIA VISA APPLICATION REVIEW")
    print("=" * 80)
    
    # Find all PDF and DOCX files
    pdf_files = list(base_path.glob("*.pdf")) + list(base_path.glob("*/*.pdf"))
    docx_files = list(base_path.glob("*.docx")) + list(base_path.glob("*/*.docx"))
    
    all_files = sorted(pdf_files + docx_files)
    
    # Extract and save all documents
    output_file = "EXTRACTED_DOCUMENTS.txt"
    
    with open(output_file, "w", encoding="utf-8") as out:
        for file_path in all_files:
            out.write("\n" + "=" * 80 + "\n")
            out.write(f"FILE: {file_path.name}\n")
            out.write("=" * 80 + "\n")
            
            print(f"Extracting: {file_path.name}...", end=" ")
            
            if file_path.suffix.lower() == ".pdf":
                content = extract_pdf_text(file_path)
            elif file_path.suffix.lower() == ".docx":
                content = extract_docx_text(file_path)
            else:
                content = "Unsupported file type"
            
            out.write(content + "\n")
            print("✓ Done")
    
    print("\n" + "=" * 80)
    print(f"✓ Extraction complete! Results saved to: {output_file}")
    print("=" * 80)
    
    # Display the extracted content
    print("\nDISPLAYING EXTRACTED CONTENT:\n")
    with open(output_file, "r", encoding="utf-8") as f:
        print(f.read())

if __name__ == "__main__":
    main()
