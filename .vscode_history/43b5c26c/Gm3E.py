import os
import sys
from pathlib import Path
import json
from datetime import datetime

# Install required libraries
try:
    import PyPDF2
except ImportError:
    os.system(f"{sys.executable} -m pip install PyPDF2 -q")
    import PyPDF2

try:
    from docx import Document
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document

def extract_pdf_text(pdf_path):
    """Extract text from PDF file"""
    try:
        text = ""
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += f"\n--- PAGE {page_num + 1} ---\n"
                text += page.extract_text()
        return text
    except Exception as e:
        return f"ERROR reading {pdf_path}: {str(e)}"

def extract_docx_text(docx_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " | "
                text += "\n"
        return text
    except Exception as e:
        return f"ERROR reading {docx_path}: {str(e)}"

def parse_personal_info():
    """Extract and organize all personal information"""
    info = {
        "Md. Mohibur Rahman (Father)": {
            "Date of Birth": "31 Dec 1968",
            "National ID": "6427025512",
            "Passport Number": "A11377254",
            "Father's Name": "Rokib Ali",
            "Mother's Name": "Joytera Bibi",
            "Address": "Shoforpur, Ward No-02, Barlekha, Dakshinbag-3252, Moulvibazar",
            "Circle (Tax)": "Circle-015 (Kulaura)",
            "TIN": "24465223154"
        },
        "Tahmidur Rahman (Son)": {
            "Date of Birth": "NOT PROVIDED",
            "National ID": "NOT PROVIDED",
            "Passport Number": "NOT PROVIDED",
            "Employment Status": "NOT PROVIDED",
            "Address": "NOT PROVIDED",
            "Personal Info": "COMPLETELY MISSING"
        }
    }
    return info

def parse_financial_info():
    """Extract all financial information"""
    finance = {
        "Exchange Rates Used": {
            "Dec 15, 2025": {"rate": 80.54, "source": "Valuation Report"},
            "Jan 21, 2026": {"rate": 82.61, "source": "Affidavit"},
            "Current (Jan 24, 2026)": {"rate": 82.61, "source": "Assumed from latest affidavit"}
        },
        "Fixed Assets (Property)": {
            "Khatian No. 52": {
                "Location": "Mouza-Dakshinbag, P.S-Barlekha, Dist: Moulvibazar",
                "JL No": "130",
                "Plot No": "1254, 8270, 8312",
                "Area": "110 Decimals",
                "Land Type": "Homestead, Cara, Amon",
                "Value_BDT": 8855000,
                "Value_AUD_Dec_Rate": 160789.66,
                "Value_AUD_Jan_Rate": 107193.03,
                "Discrepancy_AUD": 53596.63
            },
            "Khatian No. 47": {
                "Location": "Mouza-Dakshinbag, P.S-Barlekha, Dist: Moulvibazar",
                "JL No": "96",
                "Plot No": "7763, 6322",
                "Area": "63 Decimals",
                "Land Type": "Amon, Cara",
                "Value_BDT": 4095000,
                "Value_AUD_Dec_Rate": 49571.48,
                "Value_AUD_Jan_Rate": 49571.48,
                "Note": "Jan rate calculation seems incorrect"
            },
            "Total Property": {
                "Value_BDT": 12950000,
                "Value_AUD_Dec_Rate": 210361.14,
                "Value_AUD_Jan_Rate": 156764.51,
                "Total_Discrepancy": 53596.63
            }
        },
        "Cash & Bank Accounts": {
            "Account 1 (Pubali Bank - Savings)": {
                "Account Type": "Saving Account",
                "Account No": "2810101017974",
                "Branch": "Pubali Bank PLC (Dakshinbagh Branch)",
                "Balance_BDT": 1045254.61,
                "Balance_AUD": 12653.19
            },
            "Account 2 (Pubali Bank - FDR)": {
                "Account Type": "Fixed Deposit Receipt (FDR)",
                "Account No": "2810104003203",
                "Branch": "Pubali Bank PLC (Dakshinbagh Branch)",
                "Balance_BDT": 1920000.00,
                "Balance_AUD": 23242.31
            },
            "Account 3 (Uttara Bank - Current)": {
                "Account Type": "Current Account",
                "Account No": "086911100117119",
                "Branch": "Uttara Bank PLC (Ambarkhana Branch)",
                "Balance_BDT": 1397285.17,
                "Balance_AUD": 16914.65
            },
            "Total Cash": {
                "Balance_BDT": 4362539.78,
                "Balance_AUD": 52810.15
            }
        },
        "Business & Rental Income": {
            "Business Income (M.R. Traders)": {
                "Description": "Annual Business Income from M/S M.R. Traders",
                "Financial Year": "2024-2025",
                "Income_BDT": 2000000,
                "Income_AUD": 24210.74,
                "Tax Status": "CRITICAL ISSUE - Tax return shows only 410,000 BDT"
            },
            "Building Rental Income (Mohibur Plaza)": {
                "Description": "Yearly Annual Income from 5-Storey Building (Mohibur Plaza)",
                "Financial Year": "2024-2025",
                "Income_BDT": 1440000,
                "Income_AUD": 17431.73
            },
            "Total Annual Income": {
                "Income_BDT": 3440000,
                "Income_AUD": 41642.47
            }
        },
        "Tax Information": {
            "Taxpayer Name": "Md Mohibur Rahman",
            "NID": "6427025512",
            "TIN": "24465223154",
            "Circle": "Circle-015 (Kulaura)",
            "Tax Zone": "Sylhet",
            "Assessment Year": "2025-2026",
            "Date Submitted": "09/09/2025",
            "Total Income Shown": 410000,
            "Total Tax Paid": 3000,
            "CRITICAL_DISCREPANCY": "Claimed income is 3.44M BDT but tax return shows only 0.41M BDT"
        }
    }
    return finance

def parse_document_dates():
    """Extract all document dates"""
    dates = {
        "Khatian No. 52": "16-10-2022",
        "Khatian No. 47": "16-10-2022",
        "Property Valuation Report": "15-12-2025",
        "Valuation Report Issued": "09-11-2025",
        "Tax Return Submitted": "09-09-2025",
        "Survey Date": "12-12-2025",
        "Survey Intimation Date": "15-12-2025",
        "Travel Itinerary Dates": "15 Feb 2026 - 28 Feb 2026",
        "Bank Statement Dates": "NOT PROVIDED - CRITICAL ISSUE",
        "Affidavit Date": "NOT PROVIDED",
        "Marriage Certificate": "NOT PROVIDED",
        "Birth Certificate (Son)": "NOT PROVIDED"
    }
    return dates

def parse_travel_itinerary():
    """Extract travel itinerary details"""
    itinerary = {
        "Trip Duration": "14 days (15 Feb - 28 Feb 2026)",
        "Expected Duration": "13 days (user requirement)",
        "Status": "MISMATCH - 14 days vs 13 days required",
        "Destination Country": "Australia",
        "Applicants": ["Md. Mohibur Rahman (Father)", "Tahmidur Rahman (Son)"],
        "Purpose": "Tourism, cultural exploration, and family bonding",
        "Stated Purpose": "NOT Eid-specific - ISSUE",
        "Cities": {
            "Sydney": {
                "Dates": "15-19 Feb 2026",
                "Days": 5,
                "Activities": [
                    "Sydney Opera House and Harbour",
                    "Circular Quay and Darling Harbour",
                    "Bondi Beach visit and coastal walk",
                    "Blue Mountains day excursion",
                    "Museum visits"
                ]
            },
            "Melbourne": {
                "Dates": "20-23 Feb 2026",
                "Days": 4,
                "Activities": [
                    "Federation Square",
                    "Great Ocean Road full-day tour",
                    "Royal Botanic Gardens",
                    "Melbourne Museum",
                    "Local markets"
                ]
            },
            "Gold Coast": {
                "Dates": "24-27 Feb 2026",
                "Days": 4,
                "Activities": [
                    "Surfers Paradise beach",
                    "Theme park (Sea World or Movie World)",
                    "Currumbin Wildlife Sanctuary",
                    "Leisure activities"
                ]
            },
            "Departure": {
                "Date": "28 Feb 2026",
                "Status": "Return to Bangladesh"
            }
        },
        "Accommodation": "Pre-booked hotels (NO CONFIRMATION PROVIDED)",
        "Transportation": [
            "Domestic flights between cities",
            "Public transit and taxis"
        ],
        "Financial Coverage": "Fully covered by Md. Mohibur Rahman",
        "Return Assurance": "Strong ties to Bangladesh ensure return"
    }
    return itinerary

def generate_comprehensive_report():
    """Generate comprehensive analysis"""
    
    base_path = Path(".")
    output_file = "COMPREHENSIVE_DEEP_REVIEW.txt"
    
    personal = parse_personal_info()
    financial = parse_financial_info()
    dates = parse_document_dates()
    itinerary = parse_travel_itinerary()
    
    with open(output_file, "w", encoding="utf-8") as f:
        f.write("=" * 100 + "\n")
        f.write("AUSTRALIA VISA APPLICATION - COMPREHENSIVE DEEP REVIEW\n")
        f.write("Date of Analysis: " + datetime.now().strftime("%B %d, %Y - %H:%M:%S") + "\n")
        f.write("=" * 100 + "\n\n")
        
        # SECTION 1: PERSONAL INFORMATION
        f.write("\n" + "=" * 100 + "\n")
        f.write("SECTION 1: PERSONAL INFORMATION EXTRACTED\n")
        f.write("=" * 100 + "\n\n")
        for person, details in personal.items():
            f.write(f"\n{person}:\n")
            f.write("-" * 80 + "\n")
            for key, value in details.items():
                f.write(f"  {key:.<40} {value}\n")
        
        # SECTION 2: FINANCIAL INFORMATION
        f.write("\n" + "=" * 100 + "\n")
        f.write("SECTION 2: COMPLETE FINANCIAL BREAKDOWN\n")
        f.write("=" * 100 + "\n\n")
        
        f.write("EXCHANGE RATES USED IN DOCUMENTS:\n")
        f.write("-" * 80 + "\n")
        for period, data in financial["Exchange Rates Used"].items():
            f.write(f"  {period:.<40} {data['rate']} BDT/AUD (Source: {data['source']})\n")
        
        f.write("\n\nPROPERTY ASSETS:\n")
        f.write("-" * 80 + "\n")
        for prop_name, prop_data in financial["Fixed Assets (Property)"].items():
            f.write(f"\n{prop_name}:\n")
            for key, value in prop_data.items():
                f.write(f"    {key:.<35} {value}\n")
        
        f.write("\n\nBANK ACCOUNTS:\n")
        f.write("-" * 80 + "\n")
        for acc_name, acc_data in financial["Cash & Bank Accounts"].items():
            f.write(f"\n{acc_name}:\n")
            for key, value in acc_data.items():
                f.write(f"    {key:.<35} {value}\n")
        
        f.write("\n\nBUSINESS & RENTAL INCOME:\n")
        f.write("-" * 80 + "\n")
        for inc_name, inc_data in financial["Business & Rental Income"].items():
            f.write(f"\n{inc_name}:\n")
            for key, value in inc_data.items():
                f.write(f"    {key:.<35} {value}\n")
        
        f.write("\n\nTAX INFORMATION:\n")
        f.write("-" * 80 + "\n")
        for key, value in financial["Tax Information"].items():
            f.write(f"  {key:.<40} {value}\n")
        
        # SECTION 3: DOCUMENT DATES
        f.write("\n" + "=" * 100 + "\n")
        f.write("SECTION 3: DOCUMENT DATES & TIMELINE\n")
        f.write("=" * 100 + "\n\n")
        for doc, date in dates.items():
            f.write(f"  {doc:.<50} {date}\n")
        
        # SECTION 4: TRAVEL ITINERARY
        f.write("\n" + "=" * 100 + "\n")
        f.write("SECTION 4: TRAVEL ITINERARY DETAILS\n")
        f.write("=" * 100 + "\n\n")
        f.write(f"Trip Duration: {itinerary['Trip Duration']}\n")
        f.write(f"Expected Duration: {itinerary['Expected Duration']}\n")
        f.write(f"Destination: {itinerary['Destination Country']}\n")
        f.write(f"Purpose: {itinerary['Purpose']}\n")
        f.write(f"Purpose Status: {itinerary['Stated Purpose']}\n")
        
        f.write("\nDetailed Itinerary:\n")
        f.write("-" * 80 + "\n")
        for city, details in itinerary["Cities"].items():
            if city != "Departure":
                f.write(f"\n{city.upper()}:\n")
                f.write(f"  Dates: {details['Dates']}\n")
                f.write(f"  Duration: {details['Days']} days\n")
                f.write(f"  Activities:\n")
                for activity in details['Activities']:
                    f.write(f"    - {activity}\n")
            else:
                f.write(f"\n{city.upper()}:\n")
                f.write(f"  Date: {details['Date']}\n")
                f.write(f"  Status: {details['Status']}\n")
        
        f.write(f"\nAccommodation: {itinerary['Accommodation']}\n")
        f.write(f"Transportation: {', '.join(itinerary['Transportation'])}\n")
        f.write(f"Financial Coverage: {itinerary['Financial Coverage']}\n")
        f.write(f"Return Assurance: {itinerary['Return Assurance']}\n")
        
    print(f"✓ Comprehensive report saved to: {output_file}")

def main():
    print("=" * 100)
    print("PERFORMING DEEP AND EXTENSIVE REVIEW")
    print("=" * 100 + "\n")
    
    generate_comprehensive_report()
    
    # Display the file
    print("\nDISPLAYING COMPREHENSIVE REPORT:\n")
    with open("COMPREHENSIVE_DEEP_REVIEW.txt", "r", encoding="utf-8") as f:
        print(f.read())

if __name__ == "__main__":
    main()
